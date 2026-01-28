import os
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument


class DocumentLoader:

    def __init__(self, folder_path: str):
        self.folder_path = folder_path


    def load_txt(self, path: str) -> str:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()


    def load_pdf(self, path: str) -> str:
        reader = PdfReader(path)
        text = ""

        for page in reader.pages:
            text += page.extract_text() or ""

        return text


    def load_docx(self, path: str) -> str:
        doc = DocxDocument(path)
        return "\n".join([p.text for p in doc.paragraphs])


    # ✅ NEW: attach metadata
    def load(self) -> List[Document]:

        documents = []

        for file in os.listdir(self.folder_path):

            path = os.path.join(self.folder_path, file)
            suffix = Path(file).suffix.lower()

            text = ""

            try:
                if suffix == ".txt":
                    text = self.load_txt(path)

                elif suffix == ".pdf":
                    text = self.load_pdf(path)

                elif suffix == ".docx":
                    text = self.load_docx(path)

                else:
                    continue

                documents.append(
                    Document(
                        page_content=text,
                        metadata={"source": file}   # ⭐ CRITICAL
                    )
                )

            except Exception as e:
                print(f"Failed loading {file}: {e}")

        return documents
